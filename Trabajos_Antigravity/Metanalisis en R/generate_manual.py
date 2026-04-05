import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import time
import io

BASE_URL = "https://www.dsquintana.blog"
TAG_URL = "https://www.dsquintana.blog/tag/meta-analysis/"

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def get_soup(url):
    headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64)'}
    try:
        r = requests.get(url, headers=headers, timeout=15)
        if r.status_code == 200:
            return BeautifulSoup(r.content, 'html.parser')
    except Exception as e:
        print(f"Error fetching {url}: {e}")
    return None

def extract_article_links(soup):
    links = []
    visited = set()
    for a in soup.find_all('a'):
        href = a.get('href')
        if not href: continue
        full_url = href if href.startswith('http') else BASE_URL + href
        if full_url.startswith(BASE_URL) and len(full_url) > len(BASE_URL)+2 and not any(x in full_url for x in ['/tag/', '/author/', '/page/', '/about/', '/portal/']):
            if full_url not in visited:
                title = a.get_text().strip()
                if len(title) > 5:
                    links.append({'title': title, 'url': full_url})
                    visited.add(full_url)
    
    unique_links = []
    seen = set()
    for item in links:
        if item['url'] not in seen:
            unique_links.append(item)
            seen.add(item['url'])
    return unique_links

def add_image_to_doc(doc, img_url):
    try:
        if not img_url.startswith('http'):
            if img_url.startswith('//'):
                img_url = 'https:' + img_url
            elif img_url.startswith('/'):
                img_url = BASE_URL + img_url
            else:
                return
        r = requests.get(img_url, stream=True, timeout=10)
        if r.status_code == 200:
            image_stream = io.BytesIO(r.content)
            doc.add_picture(image_stream, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

def main():
    print("🚀 Fetching index...")
    soup = get_soup(TAG_URL)
    if not soup:
        print("Failed to get index.")
        return
        
    articles = extract_article_links(soup)
    if not articles:
        print("No articles found on index, using known fallback list.")
        articles = [
            {'title': 'Three common mistakes in meta-analysis', 'url': 'https://www.dsquintana.blog/three-common-mistakes-in-meta-analysis-publication-bias-funnel-plot/'},
            {'title': 'How to identify and deal with outliers in meta-analysis', 'url': 'https://www.dsquintana.blog/how-do-you-decide-which-studies-in-a-meta-analysis-are-influential-and-should-be-removed/'},
            {'title': 'How to perform a Bayesian meta-analysis in R', 'url': 'https://www.dsquintana.blog/how-to-perform-a-bayesian-meta-analysis-in-r/'},
            {'title': 'Introducing the {metameta} R package', 'url': 'https://www.dsquintana.blog/metameta-r-package-meta-analysis/'},
            {'title': 'Equivalence-enhanced forest plots with power tiles for meta-analysis', 'url': 'https://www.dsquintana.blog/equivalence-enhanced-forest-plots-with-power-tiles/'},
            {'title': 'Sunset funnel plots for meta-analysis', 'url': 'https://www.dsquintana.blog/meta-analysis-power-plot/'},
            {'title': 'Oh my GOSH: Calculating all possible meta-analysis study combinations', 'url': 'https://www.dsquintana.blog/combinatorial-meta-analysis/'},
            {'title': 'How to make a free website in R', 'url': 'https://www.dsquintana.blog/free-website-in-r-easy/'},
            {'title': 'Labeling boxplot outliers in R', 'url': 'https://www.dsquintana.blog/labeling-boxplot-outliers/'},
            {'title': 'Equivalence testing: What to do with a non-significant result', 'url': 'https://www.dsquintana.blog/equivalence-testing-non-significant-result-p-value/'},
            {'title': 'Reproducible and future-proof analysis in R', 'url': 'https://www.dsquintana.blog/reproducible-and-future-proof-analysis-in-r-open-science/'},
            {'title': 'How to visualise research trends in R', 'url': 'https://www.dsquintana.blog/how-to-visualise-research-trends-in-r/'}
        ]
        
    print(f"✅ Found {len(articles)} articles.")

    doc = Document()
    
    # Enable Page Numbers in Footer
    section = doc.sections[0]
    footer = section.footer
    p_footer = footer.paragraphs[0]
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_footer = p_footer.add_run("Página ")
    add_page_number(run_footer)

    # Title Page
    title = doc.add_heading('Manual Profesional de Metanálisis en R', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Basado en el blog de Dan Quintana').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Table of Contents (Simulated)
    doc.add_heading('Tabla de Contenidos', level=1)
    for i, art in enumerate(articles, 1):
        doc.add_paragraph(f"{i}. {art['title']}", style='List Number')
    doc.add_page_break()

    # Process Articles
    for i, tutorial in enumerate(articles, 1):
        print(f"[{i}/{len(articles)}] Extracting: {tutorial['title']}")
        doc.add_heading(tutorial['title'], level=1)
        doc.add_paragraph(f"Fuente: {tutorial['url']}").italic = True
        
        t_soup = get_soup(tutorial['url'])
        if t_soup:
            content_div = t_soup.find('div', class_='gh-content') or t_soup.find('section', class_='post-content') or t_soup.find('div', class_='post-full-content') or t_soup.find('div', class_='c-content') or t_soup.find('article') or t_soup.find('main')
            
            if content_div:
                tags = content_div.find_all(['h2', 'h3', 'p', 'pre', 'ul', 'ol', 'code', 'img', 'blockquote'])
                for tag in tags:
                    if tag.name == 'img':
                        src = tag.get('src') or tag.get('srcset')
                        if src:
                            src = src.split(',')[0].split(' ')[0]
                            add_image_to_doc(doc, src)
                        continue
                    
                    text = tag.get_text().strip()
                    if not text: continue
                    
                    if tag.name == 'h2': doc.add_heading(text, level=2)
                    elif tag.name == 'h3': doc.add_heading(text, level=3)
                    elif tag.name == 'blockquote':
                        p = doc.add_paragraph(text)
                        p.style = 'Quote'
                    elif tag.name == 'p': 
                        doc.add_paragraph(text)
                    elif tag.name in ['pre', 'code']:
                        p_code = doc.add_paragraph(text)
                        for r in p_code.runs:
                            r.font.name = 'Courier New'
                            r.font.size = Pt(9)
                            r.font.color.rgb = RGBColor(0, 50, 100)
                    elif tag.name in ['ul', 'ol']:
                        for li in tag.find_all('li'):
                            doc.add_paragraph(li.get_text().strip(), style='List Bullet')
        doc.add_page_break()

    output_path = r"c:\Trabajos Antigravity\Metanalisis en R\Manual_Metanalisis_Profesional_v3.docx"
    doc.save(output_path)
    print(f"\n🎉 Saved to {output_path}")

if __name__ == "__main__":
    main()
