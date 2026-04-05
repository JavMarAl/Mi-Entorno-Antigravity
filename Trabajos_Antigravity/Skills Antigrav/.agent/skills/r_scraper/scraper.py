import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time
import random
import os
import io

# --- CONFIGURACIÓN MAESTRA ---
INDEX_URLS = [
    "https://statisticsglobe.com/r-programming-language",
    "https://statisticsglobe.com/r-functions-list/",
    "https://statisticsglobe.com/graphics-in-r",
    "https://statisticsglobe.com/errors-warnings-r"
]

HEADERS_LIST = [
    {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'},
    {'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/17.2 Safari/605.1.15'},
    {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64; rv:109.0) Gecko/20100101 Firefox/115.0'}
]

BLACKLIST_KEYWORDS = [
    "consulting", "courses", "tutorials", "about", "privacy policy", 
    "legal notice", "data protection", "contact", "joachim schork", 
    "statistics globe", "found here", "click for more info", "login",
    "register", "search", "category", "author", "tag", "archive",
    "advertisement", "related article", "privacy", "cookie", "newsletter"
]

def get_response(url, retries=3, stream=False):
    for i in range(retries):
        try:
            headers = random.choice(HEADERS_LIST)
            time.sleep(random.uniform(0.8, 1.8)) 
            response = requests.get(url, headers=headers, timeout=30, stream=stream)
            
            if response.status_code == 200:
                return response
            elif response.status_code == 503:
                wait = (i + 1) * 10
                print(f"   ⚠️  Bloqueo 503. Esperando {wait}s...")
                time.sleep(wait)
            else:
                return None
        except Exception as e:
            print(f"   ❌ Error: {str(e)[:50]}")
            time.sleep(5)
    return None

def get_soup(url, retries=3):
    resp = get_response(url, retries=retries)
    if resp:
        return BeautifulSoup(resp.content, 'html.parser')
    return None

def is_valid_tutorial(text, href, visited):
    if not href or 'statisticsglobe.com' not in href or href in visited:
        return False
    if any(href.lower().endswith(ext) for ext in ['.pdf', '.zip', '.txt']):
        return False
    
    t = text.lower()
    if not t or len(t) < 5:
        return False
    
    for kw in BLACKLIST_KEYWORDS:
        if kw in t:
            return False
    return True

def extract_content_tags(soup):
    article = soup.find('div', class_='entry-content')
    if not article:
        article = soup.find('div', class_='l-content')
    if not article:
        article = soup.find('main', id='page-content')
    
    if not article:
        return []
        
    return article.find_all(['h2', 'h3', 'p', 'pre', 'ul', 'ol', 'code', 'img', 'table'])

def add_image_to_doc(doc, img_url):
    try:
        if not img_url.startswith('http'):
            if img_url.startswith('//'):
                img_url = 'https:' + img_url
            elif img_url.startswith('/'):
                img_url = 'https://statisticsglobe.com' + img_url
            else:
                return
        
        resp = get_response(img_url, stream=True)
        if resp and len(resp.content) > 0:
            image_stream = io.BytesIO(resp.content)
            doc.add_picture(image_stream, width=Inches(6.0))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        pass  # Silenciar errores de imagen

def add_table_to_doc(doc, table_tag):
    try:
        rows = table_tag.find_all('tr')
        if not rows: return
        
        max_cols = 0
        for row in rows:
            cols = row.find_all(['td', 'th'])
            max_cols = max(max_cols, len(cols))
        
        if max_cols == 0: return
        
        table = doc.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        for i, row in enumerate(rows):
            cells = row.find_all(['td', 'th'])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    table.cell(i, j).text = cell.get_text().strip()
    except Exception:
        pass  # Silenciar errores de tabla

def main():
    print(">>> 🚀 Manual de R - Versión DEFINITIVA (v6.0 - Template Match)")
    print(">>> Generando manual según el molde 'Learn R Programming.docx'...\n")
    
    doc = Document()
    doc_index = Document()
    
    # Título principal (como en el molde)
    title = doc.add_heading('Learn R Programming', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Tutorial \u0026 Examples').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Fuente: Statistics Globe (https://statisticsglobe.com)').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    doc_index.add_heading('Índice', 0)

    links_to_crawl = []
    visited_urls = set()

    print("📚 Fase 1: Recolectando tutoriales de todos los índices...")
    for base in INDEX_URLS:
        print(f"   • {base}")
        soup = get_soup(base)
        if not soup: continue
        
        main_content = soup.find('div', class_='entry-content') or soup.find('div', class_='l-content') or soup.find('main')
        if main_content:
            for a in main_content.find_all('a'):
                href = a.get('href')
                text = a.get_text().strip()
                if is_valid_tutorial(text, href, visited_urls):
                    links_to_crawl.append({'title': text, 'url': href})
                    visited_urls.add(href)

    total = len(links_to_crawl)
    print(f"\n✅ {total} tutoriales identificados.\n")
    print("📖 Fase 2: Descargando contenido completo (esto tomará tiempo)...\n")

    if total == 0:
        print("❌ No se encontró contenido. Saliendo.")
        return

    for i, tutorial in enumerate(links_to_crawl, 1):
        print(f"[{i}/{total}] {tutorial['title']}")
        
        # ÍNDICE
        doc_index.add_paragraph(f"{i}. {tutorial['title']}")
        
        # MANUAL - Título del tutorial (Heading 1 como en el molde)
        doc.add_heading(tutorial['title'], level=1)
        
        # Fuente en cursiva (como en el molde)
        p_fuente = doc.add_paragraph()
        run_fuente = p_fuente.add_run(f"Fuente: {tutorial['url']}")
        run_fuente.italic = True
        
        # Extraer contenido
        t_soup = get_soup(tutorial['url'])
        if t_soup:
            tags = extract_content_tags(t_soup)
            if tags:
                for tag in tags:
                    # Imágenes
                    if tag.name == 'img':
                        src = tag.get('src') or tag.get('data-src') or tag.get('data-lazy-src')
                        if src:
                            add_image_to_doc(doc, src)
                        continue
                    
                    # Tablas
                    if tag.name == 'table':
                        add_table_to_doc(doc, tag)
                        continue

                    content_text = tag.get_text().strip()
                    if not content_text: continue
                    
                    # Filtrado de ruido
                    if any(kw in content_text.lower() for kw in ["advertisement", "related article", "newsletter", "subscribe"]):
                        continue

                    # Contenido según tipo
                    if tag.name == 'h2': 
                        doc.add_heading(content_text, level=2)
                    elif tag.name == 'h3': 
                        doc.add_heading(content_text, level=3)
                    elif tag.name == 'p': 
                        doc.add_paragraph(content_text)
                    elif tag.name in ['pre', 'code']:
                        # Código formateado (como en el molde)
                        p_code = doc.add_paragraph(content_text)
                        for run in p_code.runs:
                            run.font.name = 'Courier New'
                            run.font.size = Pt(9)
                            run.font.color.rgb = RGBColor(0, 50, 100)
                    elif tag.name in ['ul', 'ol']:
                        for li in tag.find_all('li'):
                            doc.add_paragraph(li.get_text().strip(), style='List Bullet')
            else:
                doc.add_paragraph("[Contenido no disponible]")
        else:
            doc.add_paragraph("[Error al acceder al tutorial]")
            
        doc.add_page_break()
        
        # Guardar cada 20 tutoriales
        if i % 20 == 0: 
            doc.save('Manual_R_Definitivo.docx')
            print(f"      💾 Progreso guardado: {i}/{total}")

    # Guardado final
    doc.save('Manual_R_Definitivo.docx')
    doc_index.save('Indice_R_Definitivo.docx')
    
    print(f"\n🎉 ¡COMPLETADO!")
    print(f"📄 Manual: Manual_R_Definitivo.docx ({total} tutoriales)")
    print(f"📋 Índice: Indice_R_Definitivo.docx")

if __name__ == "__main__":
    main()